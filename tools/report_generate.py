
from __future__ import annotations

import json
import logging
import re
import tempfile
from datetime import datetime
from html import escape as html_escape
from pathlib import Path

from config import ReportGenerateError
from core.data_context import DataContext
from models.report_content import ReportContent, ReportTable

logger = logging.getLogger(__name__)

_MAX_DATA_KEY_ROWS = 50


async def report_generate(report_json: str, data_context: DataContext | None = None) -> str:
    """生成格式化的报告文件。

    Args:
        report_json: 报告内容的JSON字符串，格式需符合ReportContent模型。
        data_context: 可选的DataContext，用于通过data_key嵌入DataFrame。
    """
    # json.loads fallback: LLM 可能将 dict 参数序列化为 JSON 字符串
    try:
        data = json.loads(report_json) if isinstance(report_json, str) else report_json
    except json.JSONDecodeError as e:
        raise ReportGenerateError("unknown", detail=f"报告JSON解析失败: {e}") from e

    try:
        report = ReportContent.model_validate(data)
    except Exception as e:
        raise ReportGenerateError("unknown", detail=f"报告内容校验失败: {e}") from e

    _resolve_data_keys(report, data_context)
    await _resolve_chart_keys(report, data_context)

    format_type = report.format_type.lower()
    if format_type not in ("word", "ppt", "html", "pdf"):
        raise ReportGenerateError(format_type, detail="不支持的格式，可选: word, ppt, html, pdf")

    output_dir = Path(tempfile.gettempdir()) / "bi_reports"
    output_dir.mkdir(exist_ok=True)

    try:
        if format_type == "html":
            return _generate_html(report, output_dir)
        elif format_type == "word":
            return _generate_word(report, output_dir)
        elif format_type == "ppt":
            from tools.ppt_renderer import render_ppt

            return await render_ppt(report, output_dir, data_context)
        else:  # pdf
            from tools.ppt_renderer import render_pdf

            return await render_pdf(report, output_dir)
    except ReportGenerateError:
        raise
    except Exception as e:
        raise ReportGenerateError(format_type, detail=str(e)) from e


def _resolve_data_keys(report: ReportContent, data_context: DataContext | None) -> None:
    """对每个含 data_key 的 section，从 DataContext 读取 DataFrame 转为 ReportTable。"""
    if data_context is None:
        return
    for section in report.sections:
        if not section.data_key:
            continue
        df = data_context.get(section.data_key)
        if df is None:
            logger.warning("data_key '%s' 在DataContext中不存在，跳过", section.data_key)
            continue
        headers = [str(h) for h in df.columns.tolist()]
        rows = [[str(cell) for cell in row] for row in df.head(_MAX_DATA_KEY_ROWS).itertuples(index=False)]
        caption = f"数据来源: {section.data_key}"
        if len(df) > _MAX_DATA_KEY_ROWS:
            caption += f" (共{len(df)}行，显示前{_MAX_DATA_KEY_ROWS}行)"
        section.tables.append(ReportTable(headers=headers, rows=rows, caption=caption))


async def _resolve_chart_keys(report: ReportContent, data_context: DataContext | None) -> None:
    """对每个含 chart_keys 的 section，将图表渲染为 PNG（如尚未渲染）。"""
    if data_context is None:
        return
    for section in report.sections:
        if not section.chart_keys:
            continue
        for chart_key in section.chart_keys:
            artifact = data_context.get_chart(chart_key)
            if artifact is None:
                logger.warning("chart_key '%s' 不存在，跳过", chart_key)
                continue
            # 如果已有 PNG 则跳过
            if artifact.png_path and Path(artifact.png_path).exists():
                continue
            # 用 Playwright 截图
            try:
                png_path = await _render_chart_to_png(artifact)
                artifact.png_path = str(png_path)
            except Exception as e:
                logger.warning("图表截图失败 '%s': %s", chart_key, e)


async def _render_chart_to_png(artifact: object) -> Path:
    """用 Playwright 将图表 HTML 截图为 PNG。"""
    from models.chart_artifact import ChartArtifact

    art = artifact  # type: ignore[attr-defined]
    html_path = art.html_path  # type: ignore[attr-defined]
    png_path = Path(html_path).with_suffix(".png")

    from playwright.async_api import async_playwright

    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page(viewport={"width": art.width, "height": art.height})  # type: ignore[attr-defined]
        file_url = Path(html_path).as_uri()
        await page.goto(file_url)
        await page.wait_for_load_state("networkidle")
        await page.screenshot(path=str(png_path), full_page=False)
        await browser.close()

    return png_path


def _safe_filename(title: str) -> str:
    """生成安全的文件名，移除特殊字符。"""
    safe = re.sub(r'[\\/:*?"<>|]', "_", title).strip()
    if not safe:
        safe = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return safe


# ---------------------------------------------------------------------------
# Word 渲染器
# ---------------------------------------------------------------------------


def _generate_word(report: ReportContent, output_dir: Path) -> str:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as e:
        raise ReportGenerateError("word", detail="python-docx not installed, run: uv add python-docx") from e

    doc = Document()

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # type: ignore[attr-defined]
    style.font.size = Pt(11)

    # 标题
    doc.add_heading(report.title, level=0)

    # 章节
    for section in report.sections:
        if section.heading:
            doc.add_heading(section.heading, level=section.level)
        for para in section.paragraphs:
            doc.add_paragraph(para)
        for table in section.tables:
            _render_word_table(doc, table)
        # 嵌入图表
        for chart_key in section.chart_keys:
            if data_context:
                artifact = data_context.get_chart(chart_key)
                if artifact and artifact.png_path and Path(artifact.png_path).exists():
                    from docx.shared import Inches
                    doc.add_picture(artifact.png_path, width=Inches(5.5))

    # 结论
    if report.conclusion:
        doc.add_heading("结论", level=1)
        doc.add_paragraph(report.conclusion)

    filename = f"{_safe_filename(report.title)}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    return f"Word报告已生成: {output_path}"


def _render_word_table(doc: object, table: ReportTable) -> None:
    from docx.shared import Pt

    num_cols = len(table.headers)
    if num_cols == 0:
        return
    num_rows = len(table.rows) + 1  # +1 for header
    t = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")  # type: ignore[attr-defined]

    # 表头
    for i, header in enumerate(table.headers):
        cell = t.rows[0].cells[i]  # type: ignore[index]
        cell.text = header
        for paragraph in cell.paragraphs:  # type: ignore[attr-defined]
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)  # type: ignore[attr-defined]

    # 数据行
    for row_idx, row_data in enumerate(table.rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = t.rows[row_idx + 1].cells[col_idx]  # type: ignore[index]
                cell.text = cell_text
                for paragraph in cell.paragraphs:  # type: ignore[attr-defined]
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # type: ignore[attr-defined]

    # 表格说明
    if table.caption:
        doc.add_paragraph(table.caption)  # type: ignore[attr-defined]


# ---------------------------------------------------------------------------
# HTML 渲染器（所有用户内容经 html_escape 转义，防止 XSS）
# ---------------------------------------------------------------------------


def _generate_html(report: ReportContent, output_dir: Path) -> str:
    safe_title = html_escape(report.title)
    parts: list[str] = []
    parts.append(f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{safe_title}</title>
    <style>
        body {{ font-family: "Microsoft YaHei", sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; color: #333; }}
        h1 {{ color: #333; border-bottom: 2px solid #0078d4; padding-bottom: 10px; }}
        h2 {{ color: #0078d4; margin-top: 30px; }}
        h3 {{ color: #555; margin-top: 20px; }}
        h4 {{ color: #666; margin-top: 15px; }}
        p {{ line-height: 1.8; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th {{ background-color: #0078d4; color: white; padding: 8px 12px; text-align: left; }}
        td {{ padding: 8px 12px; border: 1px solid #ddd; }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        .caption {{ font-size: 0.9em; color: #666; margin-top: 4px; }}
        .conclusion {{ margin-top: 30px; padding: 15px; background-color: #f0f7ff; border-left: 4px solid #0078d4; }}
    </style>
</head>
<body>
    <h1>{safe_title}</h1>""")

    for section in report.sections:
        # level=1 -> h2, level=2 -> h3, level=3 -> h4
        tag = f"h{section.level + 1}"
        if section.heading:
            parts.append(f"    <{tag}>{html_escape(section.heading)}</{tag}>")
        for para in section.paragraphs:
            parts.append(f"    <p>{html_escape(para)}</p>")
        for table in section.tables:
            parts.append(_render_html_table(table))
        # 嵌入图表
        for chart_key in section.chart_keys:
            if data_context:
                artifact = data_context.get_chart(chart_key)
                if artifact:
                    safe_src = html_escape(artifact.html_path)
                    parts.append(f'    <iframe src="{safe_src}" width="{artifact.width}" height="{artifact.height}" frameborder="0" style="border:1px solid #e0e0e0;border-radius:8px;margin:10px 0"></iframe>')

    if report.conclusion:
        parts.append(f'    <div class="conclusion"><h2>结论</h2><p>{html_escape(report.conclusion)}</p></div>')

    parts.append("</body>\n</html>")

    filename = f"{_safe_filename(report.title)}.html"
    output_path = output_dir / filename
    output_path.write_text("\n".join(parts), encoding="utf-8")
    return f"HTML报告已生成: {output_path}"


def _render_html_table(table: ReportTable) -> str:
    parts: list[str] = []
    parts.append("    <table>")
    parts.append("        <thead><tr>")
    for h in table.headers:
        parts.append(f"            <th>{html_escape(str(h))}</th>")
    parts.append("        </tr></thead>")
    parts.append("        <tbody>")
    for row in table.rows:
        parts.append("            <tr>")
        for cell in row:
            parts.append(f"                <td>{html_escape(str(cell))}</td>")
        parts.append("            </tr>")
    parts.append("        </tbody>")
    parts.append("    </table>")
    if table.caption:
        parts.append(f'    <p class="caption">{html_escape(table.caption)}</p>')
    return "\n".join(parts)